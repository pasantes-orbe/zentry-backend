from datetime import date
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_title(document: Document, text: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(20)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_h1(document: Document, text: str) -> None:
    document.add_heading(text, level=1)


def add_h2(document: Document, text: str) -> None:
    document.add_heading(text, level=2)


def add_para(document: Document, text: str) -> None:
    document.add_paragraph(text)


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style='List Bullet')


doc = Document()

# Title
add_title(doc, "Propuesta de Arquitectura MVP SaaS – Zentry")
add_para(doc, f"Fecha: {date.today().strftime('%Y-%m-%d')}")

# Resumen Ejecutivo
add_h1(doc, "Resumen ejecutivo")
add_para(doc, "Recomendación: Monolito modular multi-tenant con Base de Datos Shared Schema (columna 'tenant_id'), manteniendo el stack actual (Node 18, Express, Sequelize, TypeScript). Objetivo noviembre: Aislamiento de datos por tenant, subdominios por tenant, auth multi-tenant, despliegue productivo con observabilidad y rollback. Motivo: menor riesgo y tiempo (3–4 semanas) que microservicios, con camino claro de evolución.")
add_bullet(doc, "Arquitectura MVP SaaS: Monolito Modular Multi-tenant (diseño hexagonal)")
add_bullet(doc, "Base de Datos: Shared Schema con 'tenant_id' e índices compuestos")
add_bullet(doc, "Frontend: Angular 18 se mantiene sin cambios de framework")
add_bullet(doc, "Entrega noviembre: listo para vender con aislamiento por tenant")

# Alcance
add_h1(doc, "Alcance")
add_bullet(doc, "Backend multi-tenant (middleware de tenant, repositorios con scoping, auth/autorización)")
add_bullet(doc, "Migraciones de BD: agregar 'tenant_id', índices y unicidad por tenant a tablas de negocio")
add_bullet(doc, "Subdominios {tenant}.zentry.com y/o header seguro X-Tenant-Id desde el gateway")
add_bullet(doc, "CI/CD, logging/metrics por tenant, rate limiting, backups y runbook")
add_bullet(doc, "No incluye: cambio de framework frontend ni migración a microservicios en esta fase")

# Arquitectura técnica
add_h1(doc, "Arquitectura técnica")
add_h2(doc, "Estilo")
add_para(doc, "Monolito modular con principios de Arquitectura Hexagonal/Clean Architecture. Capas: Dominio (reglas por módulo), Aplicación (casos de uso), Infraestructura (Express, Sequelize, adaptadores). Límites bien definidos para futura extracción por 'strangler pattern'.")
add_h2(doc, "Modelo multi-tenant")
add_bullet(doc, "DB Shared Schema: columna obligatoria 'tenant_id' en todas las tablas de negocio y tablas de unión relacionadas")
add_bullet(doc, "Índices y unicidad: índices compuestos que incluyan 'tenant_id'; convertir UNIQUE(col) a UNIQUE(tenant_id, col)")
add_bullet(doc, "Integridad: FKs alineadas al tenant; validación en la capa de dominio para evitar cruces")
add_bullet(doc, "Multi-country (si aplica): agregar 'country_code' como dimensión; índices por (tenant_id, country_code, …)")
add_h2(doc, "Resolución de tenant en runtime")
add_bullet(doc, "Subdominio {tenant}.zentry.com como fuente principal; fallback por header X-Tenant-Id")
add_bullet(doc, "Middleware: resuelve tenantId al inicio y lo inyecta en un contexto por request")
add_bullet(doc, "Repositorios/ORM: scoping automático 'WHERE tenant_id = <contexto>' en lecturas/escrituras")
add_h2(doc, "Autenticación y autorización")
add_bullet(doc, "JWT con tenant_id, user_id y roles/claims")
add_bullet(doc, "Policies/guards por módulo verificando pertenencia al tenant y rol")
add_h2(doc, "Observabilidad y seguridad")
add_bullet(doc, "Logging/metrics trazadas por tenantId (y country si aplica)")
add_bullet(doc, "Rate limiting por tenant e IP; auditoría de acciones con tenant_id")
add_bullet(doc, "Backups diarios y pruebas de restore; secretos por entorno")

# Despliegue y lanzamiento
add_h1(doc, "Despliegue y lanzamiento")
add_h2(doc, "Infraestructura")
add_bullet(doc, "Contenedores Docker (Node 18) para el backend monolítico; Nginx/Ingress como gateway")
add_bullet(doc, "Entornos: dev, staging, prod con mismas imágenes y distinta configuración")
add_bullet(doc, "CI/CD: pipeline con tests, migraciones automáticas, despliegue y healthchecks")
add_h2(doc, "Dominios y TLS")
add_bullet(doc, "DNS wildcard *.zentry.com apuntando al gateway")
add_bullet(doc, "SSL wildcard con renovación automática")
add_bullet(doc, "Enrutamiento por subdominio hacia backend; inyección de X-Tenant-Id si se requiere")
add_h2(doc, "Base de datos")
add_bullet(doc, "Migraciones: agregar tenant_id, índices y unicidades por tenant a las 18 tablas priorizadas")
add_bullet(doc, "Backfill: asignación de tenant_id a datos existentes con criterio definido")
add_bullet(doc, "Seeds: creación de tenants y planes/feature flags básicos")
add_h2(doc, "Plan de rollout")
add_bullet(doc, "Canary/piloto con 1–2 tenants")
add_bullet(doc, "Smoke tests de endpoints críticos por tenant")
add_bullet(doc, "Feature flags por tenant/plan si aplica; estrategia de rollback clara")

# Cronograma
add_h1(doc, "Cronograma estimado (3–4 semanas)")
add_h2(doc, "Semana 1")
add_bullet(doc, "Clasificación de tablas (tenant-owned/global) y migraciones base: tenant_id, índices y unicidad por tenant")
add_bullet(doc, "Middleware de tenant y contexto por request")
add_h2(doc, "Semana 2")
add_bullet(doc, "Repositorios con scoping automático; refactor de endpoints críticos")
add_bullet(doc, "Tests E2E con 2 tenants para validar aislamiento de datos")
add_h2(doc, "Semana 3")
add_bullet(doc, "JWT con tenant_id y roles; subdominios y gateway en producción")
add_bullet(doc, "Observabilidad (logs/metrics por tenant) y rate limiting por tenant")
add_h2(doc, "Semana 4")
add_bullet(doc, "Hardening, QA de regresión, optimización de índices, runbooks, canary y salida a prod")

# Criterios listo para vender
add_h1(doc, "Criterios 'listo para vender' (MVP SaaS)")
add_bullet(doc, "Aislamiento de datos por tenant validado en E2E")
add_bullet(doc, "Acceso por subdominio y autenticación multi-tenant operativos")
add_bullet(doc, "Operabilidad: logs/metrics, rate limit, backups y rollback funcionando")
add_bullet(doc, "Onboarding: creación de tenants, usuarios y roles; planes/flags básicos")

# Riesgos y mitigación
add_h1(doc, "Riesgos y mitigación")
add_bullet(doc, "Consultas cross-tenant o joins globales: cubrir con scoping obligatorio y tests E2E")
add_bullet(doc, "Unicidades globales (p.ej. email): convertir a unicidad por tenant o validar en dominio")
add_bullet(doc, "Reportes globales: filtrar explícitamente por tenant o generar vistas por tenant")

# Coste y capacidad
add_h1(doc, "Coste y capacidad")
add_bullet(doc, "Capacidad suficiente para 20–30 usuarios y 3–10 países por cliente")
add_bullet(doc, "Coste operativo bajo: un monolito escalable horizontalmente; monitoreo por tenant")

# Evolución futura
add_h1(doc, "Evolución futura (post-MVP)")
add_bullet(doc, "Schema-per-tenant si se requiere mayor aislamiento/compliance")
add_bullet(doc, "Extracción gradual a microservicios por módulos con alto throughput o ciclos propios")

# Aclaración clave
add_h1(doc, "Aclaración clave: ¿es solamente monolito?")
add_para(doc, "En despliegue, sí, es un monolito (una aplicación), lo que simplifica y acelera el MVP. En diseño, es modular/hexagonal con límites definidos; esto permite extraer servicios cuando el producto escale, sin reescribir todo. Es arquitectura SaaS desde el día 1 (multi-tenant, subdominios, auth y operaciones por tenant), aun sin microservicios.")

# Frontend
add_h1(doc, "Frontend (Angular 18)")
add_para(doc, "No se cambia de framework. El frontend sólo debe incluir el tenant en la base URL o enviar X-Tenant-Id; el trabajo principal está en backend + despliegue.")

# Save
output_path = "/workspace/Propuesta_Arquitectura_MVP_SaaS_Zentry.docx"
doc.save(output_path)
print(f"Documento generado: {output_path}")
